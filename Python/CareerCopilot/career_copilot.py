from pathlib import Path
import re
import os
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

PROJECT_PATH = Path(__file__).parent
REPORT_FILE = PROJECT_PATH / "reports" / "career_report.txt"

STOP_WORDS = {
    "we", "are", "a", "an", "the", "and", "or", "with", "in", "to",
    "of", "for", "on", "is", "be", "as", "this", "that", "you",
    "your", "our", "seeking", "experience", "experienced", "role",
    "candidate", "will", "have", "has", "from", "by", "at", "responsible"
}

IMPORTANT_KEYWORDS = {
    "salesforce", "administrator", "admin", "crm", "flow", "automation",
    "reports", "dashboards", "tableau", "excel", "sql", "soql",
    "business", "analyst", "stakeholder", "communication", "forecasting",
    "pipeline", "revops", "revenue", "operations", "data", "analytics",
    "process", "requirements", "documentation", "testing", "training",
    "integration", "workflow", "optimization", "reporting"
}

IMPORTANT_PHRASES = [
    "salesforce administrator", "business analyst", "stakeholder communication",
    "revenue operations", "sales operations", "pipeline management",
    "flow automation", "salesforce flow", "crm reporting", "data analysis",
    "business requirements", "process improvement", "executive dashboard",
    "user acceptance testing", "data quality", "salesforce reports",
    "salesforce dashboards", "workflow automation"
]

resume_path = None
job_path = None
latest_report = ""
latest_missing_keywords = []
latest_missing_phrases = []
latest_score = 0
latest_estimated_score = 0
latest_resume_suggestions = ""


def read_file(path):
    path = Path(path)
    file_type = path.suffix.lower()

    if file_type == ".txt":
        return path.read_text(encoding="utf-8")

    if file_type == ".docx":
        document = Document(path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    if file_type == ".pdf":
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    messagebox.showerror("Unsupported File", "Please upload a .txt, .docx, or .pdf file.")
    return ""


def extract_words(text):
    words = set(re.findall(r"\b[a-zA-Z0-9+#.-]+\b", text.lower()))
    return words - STOP_WORDS


def find_phrases(text):
    text = text.lower()
    return sorted([phrase for phrase in IMPORTANT_PHRASES if phrase in text])


def get_rating(score):
    if score >= 90:
        return "Excellent Match", "green"
    if score >= 75:
        return "Good Match", "orange"
    if score >= 50:
        return "Fair Match", "goldenrod"
    return "Needs Work", "red"


def get_priority_items(missing_keywords, missing_phrases):
    high_priority_terms = {
        "salesforce", "salesforce administrator", "flow automation",
        "salesforce flow", "business analyst", "stakeholder communication",
        "business requirements", "crm reporting", "data analysis",
        "reports", "dashboards", "sql", "soql"
    }

    missing_items = missing_phrases + missing_keywords
    high_priority = sorted([item for item in missing_items if item in high_priority_terms])
    lower_priority = sorted([item for item in missing_items if item not in high_priority_terms])

    return high_priority, lower_priority


def build_resume_suggestions(missing_keywords, missing_phrases):
    missing_items = missing_phrases + missing_keywords
    suggestions = []

    suggestions.append("Resume Improvement Assistant")
    suggestions.append("=" * 45)
    suggestions.append("Use only the bullets that truthfully match your experience.")
    suggestions.append("")

    suggestions.append("Missing ATS Items")
    suggestions.append("-" * 25)

    if missing_items:
        for item in missing_items:
            suggestions.append(f"• {item}")
    else:
        suggestions.append("No major missing ATS items found.")

    suggestions.append("")
    suggestions.append("Paste-Ready Resume Bullet Ideas")
    suggestions.append("-" * 35)

    if "flow automation" in missing_items or "salesforce flow" in missing_items or "automation" in missing_items:
        suggestions.append("• Built Salesforce Flow automations to reduce manual work, improve CRM process consistency, and support business workflow efficiency.")

    if "soql" in missing_items:
        suggestions.append("• Used SOQL to query Salesforce data, validate records, and support CRM reporting and analysis.")

    if "sql" in missing_items:
        suggestions.append("• Used SQL to retrieve, clean, and analyze data in support of reporting and business decision-making.")

    if "stakeholder communication" in missing_items or "stakeholder" in missing_items:
        suggestions.append("• Partnered with stakeholders to gather requirements, clarify business needs, and support solution design.")

    if "forecasting" in missing_items or "pipeline" in missing_items:
        suggestions.append("• Supported sales pipeline visibility through CRM reports, dashboards, and forecasting-related analysis.")

    if "reports" in missing_items or "dashboards" in missing_items or "crm reporting" in missing_items:
        suggestions.append("• Created CRM reports and dashboards to monitor performance, identify trends, and support decision-making.")

    if "integration" in missing_items:
        suggestions.append("• Supported CRM process improvements by documenting requirements, identifying data needs, and coordinating system workflow enhancements.")

    if len(suggestions) <= 12:
        suggestions.append("• Strengthened CRM documentation, reporting, and process improvement support across business workflows.")

    return "\n".join(suggestions)


def build_bar_chart(score):
    filled = round(score / 10)
    empty = 10 - filled
    return "█" * filled + "░" * empty


def build_report(score, rating, matched_phrases, missing_phrases, matched_keywords, missing_keywords, resume_words, job_words):
    global latest_estimated_score, latest_resume_suggestions

    high_priority, lower_priority = get_priority_items(missing_keywords, missing_phrases)
    top_missing = (high_priority + lower_priority)[:5]
    estimated_score = min(score + (len(top_missing) * 4), 100)
    latest_estimated_score = estimated_score

    resume_suggestions = build_resume_suggestions(missing_keywords, missing_phrases)
    latest_resume_suggestions = resume_suggestions

    report = []
    report.append("=" * 60)
    report.append("CAREERCOPILOT REPORT")
    report.append("=" * 60)
    report.append(f"Run Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append("")
    report.append("Recruiter Summary")
    report.append("-" * 30)
    report.append(f"Overall Match: {score}%")
    report.append(f"Rating: {rating}")
    report.append(f"ATS Visual: {build_bar_chart(score)}")
    report.append("")
    report.append("Estimated Improvement")
    report.append("-" * 30)
    report.append(f"Current Score: {score}%")
    report.append(f"Estimated Score If Legitimate Missing Items Are Added: {estimated_score}%")
    report.append("")
    report.append("Top Missing Items")
    report.append("-" * 30)

    if top_missing:
        for item in top_missing:
            report.append(f"✘ {item}")
    else:
        report.append("None found.")

    report.append("")
    report.append("ATS Strengths")
    report.append("-" * 30)

    for item in matched_phrases + matched_keywords:
        report.append(f"✔ {item}")

    if not matched_phrases and not matched_keywords:
        report.append("None found.")

    report.append("")
    report.append("Priority Missing Items")
    report.append("-" * 30)
    report.append("HIGH PRIORITY")

    if high_priority:
        for item in high_priority:
            report.append(f"✘ {item}")
    else:
        report.append("None found.")

    report.append("")
    report.append("LOWER PRIORITY")

    if lower_priority:
        for item in lower_priority:
            report.append(f"✘ {item}")
    else:
        report.append("None found.")

    report.append("")
    report.append("Quick Stats")
    report.append("-" * 30)
    report.append(f"Resume keyword count: {len(resume_words)}")
    report.append(f"Job description keyword count: {len(job_words)}")
    report.append(f"Matched keywords: {len(matched_keywords)}")
    report.append(f"Missing keywords: {len(missing_keywords)}")
    report.append(f"Matched phrases: {len(matched_phrases)}")
    report.append(f"Missing phrases: {len(missing_phrases)}")
    report.append("")
    report.append("ATS Score Scale")
    report.append("-" * 30)
    report.append("90-100   Excellent Match")
    report.append("75-89    Strong Match")
    report.append("50-74    Moderate Match")
    report.append("0-49     Needs Work")
    report.append("")
    report.append(resume_suggestions)

    return "\n".join(report)


def draw_donut(score, color):
    score_canvas.delete("all")
    score_canvas.create_oval(10, 10, 210, 210, outline="#ddd", width=18)

    extent = int((score / 100) * 359)
    score_canvas.create_arc(10, 10, 210, 210, start=90, extent=-extent, outline=color, width=18, style="arc")
    score_canvas.create_text(110, 94, text=f"{score}%", font=("Arial", 30, "bold"), fill=color)
    score_canvas.create_text(110, 130, text="ATS Match", font=("Arial", 10), fill="black")


def set_progress_color(color):
    style.configure("score.Horizontal.TProgressbar", troughcolor="#e6e6e6", background=color)


def save_report_as_pdf(save_path):
    pdf = canvas.Canvas(save_path, pagesize=letter)
    width, height = letter
    y = height - 40

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(40, y, "CareerCopilot Report")
    y -= 30

    pdf.setFont("Helvetica", 9)

    for line in latest_report.split("\n"):
        if y < 40:
            pdf.showPage()
            pdf.setFont("Helvetica", 9)
            y = height - 40

        clean_line = line.replace("✔", "[MATCH]").replace("✘", "[MISSING]").replace("█", "#").replace("░", "-")
        pdf.drawString(40, y, clean_line[:100])
        y -= 13

    pdf.save()


def save_report_as_txt(save_path):
    Path(save_path).write_text(latest_report, encoding="utf-8")


def save_report_as_docx(save_path):
    document = Document()
    document.add_heading("CareerCopilot Report", level=1)

    for line in latest_report.split("\n"):
        if line.strip() == "":
            document.add_paragraph("")
        elif line.isupper() and len(line) < 40:
            document.add_heading(line.title(), level=2)
        else:
            document.add_paragraph(line)

    document.save(save_path)


def export_report():
    if not latest_report:
        messagebox.showwarning("No Report", "Please analyze a resume first.")
        return

    save_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[
            ("PDF Files", "*.pdf"),
            ("Text Files", "*.txt"),
            ("Word Documents", "*.docx"),
        ],
        initialfile="career_report.pdf",
        title="Export Report As"
    )

    if not save_path:
        return

    file_type = Path(save_path).suffix.lower()

    if file_type == ".pdf":
        save_report_as_pdf(save_path)
    elif file_type == ".txt":
        save_report_as_txt(save_path)
    elif file_type == ".docx":
        save_report_as_docx(save_path)
    else:
        messagebox.showerror("Unsupported Export", "Please save as .pdf, .txt, or .docx.")
        return

    messagebox.showinfo("Report Exported", f"Report saved to:\n{save_path}")


def copy_report():
    if not latest_report:
        messagebox.showwarning("No Report", "Please analyze a resume first.")
        return

    app.clipboard_clear()
    app.clipboard_append(latest_report)
    messagebox.showinfo("Copied", "Full report copied to clipboard.")


def open_reports_folder():
    reports_path = PROJECT_PATH / "reports"
    reports_path.mkdir(exist_ok=True)
    os.startfile(reports_path)


def show_improve_resume_popup():
    if not latest_resume_suggestions:
        messagebox.showwarning("No Analysis", "Please analyze a resume first.")
        return

    popup = tk.Toplevel(app)
    popup.title("Resume Improvement Assistant")
    popup.geometry("760x620")

    header = tk.Label(popup, text="Resume Improvement Assistant", font=("Arial", 20, "bold"))
    header.pack(pady=10)

    score_text = f"Current ATS Score: {latest_score}%   |   Estimated Improved Score: {latest_estimated_score}%"
    score_summary = tk.Label(popup, text=score_text, font=("Arial", 12, "bold"))
    score_summary.pack(pady=5)

    text_area = scrolledtext.ScrolledText(popup, width=90, height=28)
    text_area.pack(padx=15, pady=10)
    text_area.insert(tk.END, latest_resume_suggestions)

    def copy_bullets():
        app.clipboard_clear()
        app.clipboard_append(latest_resume_suggestions)
        messagebox.showinfo("Copied", "Resume improvement suggestions copied to clipboard.")

    button_frame = tk.Frame(popup)
    button_frame.pack(pady=8)

    tk.Button(button_frame, text="📋 Copy Suggestions", width=22, command=copy_bullets).grid(row=0, column=0, padx=8)
    tk.Button(button_frame, text="Close", width=12, command=popup.destroy).grid(row=0, column=1, padx=8)


def analyze():
    global resume_path, job_path, latest_missing_keywords, latest_missing_phrases, latest_report, latest_score

    if not resume_path or not job_path:
        messagebox.showwarning("Missing Files", "Please load both a resume and a job description.")
        return

    status_label.config(text="Status: Reading files...")
    progress_bar["value"] = 20
    app.update_idletasks()

    resume_text = read_file(resume_path)
    job_text = read_file(job_path)

    status_label.config(text="Status: Extracting keywords...")
    progress_bar["value"] = 45
    app.update_idletasks()

    resume_words = extract_words(resume_text)
    job_words = extract_words(job_text)

    job_keywords = sorted(job_words & IMPORTANT_KEYWORDS)
    resume_keywords = sorted(resume_words & IMPORTANT_KEYWORDS)

    matched_keywords = sorted(set(job_keywords) & set(resume_keywords))
    missing_keywords = sorted(set(job_keywords) - set(resume_keywords))

    status_label.config(text="Status: Matching phrases...")
    progress_bar["value"] = 70
    app.update_idletasks()

    job_phrases = find_phrases(job_text)
    resume_phrases = find_phrases(resume_text)

    matched_phrases = sorted(set(job_phrases) & set(resume_phrases))
    missing_phrases = sorted(set(job_phrases) - set(resume_phrases))

    latest_missing_keywords = missing_keywords
    latest_missing_phrases = missing_phrases

    total_possible = len(job_keywords) + len(job_phrases)
    total_matched = len(matched_keywords) + len(matched_phrases)
    score = round((total_matched / total_possible) * 100) if total_possible else 0
    latest_score = score

    rating, color = get_rating(score)

    final_report = build_report(
        score,
        rating,
        matched_phrases,
        missing_phrases,
        matched_keywords,
        missing_keywords,
        resume_words,
        job_words,
    )

    latest_report = final_report

    output_box.delete("1.0", tk.END)
    output_box.insert(tk.END, final_report)

    REPORT_FILE.parent.mkdir(exist_ok=True)
    REPORT_FILE.write_text(final_report, encoding="utf-8")

    score_label.config(text=f"{score}%", fg=color)
    rating_label.config(text=rating, fg=color)
    summary_label.config(
        text=f"Matched: {len(matched_keywords)} keywords | Missing: {len(missing_keywords)} keywords | Potential Score: {latest_estimated_score}%"
    )

    draw_donut(score, color)
    set_progress_color(color)

    status_label.config(text="Status: Complete")
    progress_bar["value"] = 100

    messagebox.showinfo("Report Saved", f"TXT report automatically saved to:\n{REPORT_FILE}")


def load_resume():
    global resume_path
    resume_path = filedialog.askopenfilename(
        filetypes=[
            ("Supported Files", "*.txt *.docx *.pdf"),
            ("Text Files", "*.txt"),
            ("Word Documents", "*.docx"),
            ("PDF Files", "*.pdf"),
        ]
    )
    if resume_path:
        resume_label.config(text=f"Resume: {Path(resume_path).name}")


def load_job():
    global job_path
    job_path = filedialog.askopenfilename(
        filetypes=[
            ("Supported Files", "*.txt *.docx *.pdf"),
            ("Text Files", "*.txt"),
            ("Word Documents", "*.docx"),
            ("PDF Files", "*.pdf"),
        ]
    )
    if job_path:
        job_label.config(text=f"Job Description: {Path(job_path).name}")


def copy_ats_keywords():
    missing_items = latest_missing_phrases + latest_missing_keywords

    if not missing_items:
        messagebox.showinfo("Nothing to Copy", "No missing ATS keywords or phrases found yet.")
        return

    copied_text = "\n".join(missing_items)
    app.clipboard_clear()
    app.clipboard_append(copied_text)
    messagebox.showinfo("Copied", f"Copied {len(missing_items)} ATS keyword(s)/phrase(s) to clipboard.")


def clear_output():
    global latest_report, latest_missing_keywords, latest_missing_phrases, latest_score, latest_estimated_score, latest_resume_suggestions

    output_box.delete("1.0", tk.END)
    score_label.config(text="--%", fg="black")
    rating_label.config(text="Not analyzed yet", fg="black")
    summary_label.config(text="Analysis summary will appear here after scanning.")
    status_label.config(text="Status: Ready")
    progress_bar["value"] = 0
    set_progress_color("#bdbdbd")
    latest_report = ""
    latest_missing_keywords = []
    latest_missing_phrases = []
    latest_score = 0
    latest_estimated_score = 0
    latest_resume_suggestions = ""
    score_canvas.delete("all")
    score_canvas.create_oval(10, 10, 210, 210, outline="#ddd", width=18)
    score_canvas.create_text(110, 110, text="--%", font=("Arial", 30, "bold"), fill="black")


app = tk.Tk()
app.title("CareerCopilot")
app.geometry("1240x980")

style = ttk.Style()
style.theme_use("default")
style.configure("score.Horizontal.TProgressbar", troughcolor="#e6e6e6", background="#bdbdbd")

title = tk.Label(app, text="CareerCopilot", font=("Arial", 28, "bold"))
title.pack(pady=4)

subtitle = tk.Label(app, text="Resume + Job Description Match Analyzer", font=("Arial", 13))
subtitle.pack(pady=1)

top_frame = tk.Frame(app)
top_frame.pack(pady=4)

score_canvas = tk.Canvas(top_frame, width=220, height=220, highlightthickness=0)
score_canvas.grid(row=0, column=0, padx=24)

score_canvas.create_oval(10, 10, 210, 210, outline="#ddd", width=18)
score_canvas.create_text(110, 110, text="--%", font=("Arial", 30, "bold"), fill="black")

score_info_frame = tk.Frame(top_frame)
score_info_frame.grid(row=0, column=1, padx=24)

score_label = tk.Label(score_info_frame, text="--%", font=("Arial", 46, "bold"))
score_label.pack(pady=2)

rating_label = tk.Label(score_info_frame, text="Not analyzed yet", font=("Arial", 17, "bold"))
rating_label.pack(pady=2)

scale_label = tk.Label(score_info_frame, text="90+ Excellent | 75+ Strong | 50+ Fair | Below 50 Needs Work", font=("Arial", 10))
scale_label.pack(pady=2)

summary_label = tk.Label(score_info_frame, text="Analysis summary will appear here after scanning.", font=("Arial", 10, "bold"))
summary_label.pack(pady=8)

button_frame = tk.Frame(app)
button_frame.pack(pady=8)

tk.Button(button_frame, text="📄 Load Resume", width=17, command=load_resume).grid(row=0, column=0, padx=3, pady=4)
tk.Button(button_frame, text="💼 Load Job Description", width=23, command=load_job).grid(row=0, column=1, padx=3, pady=4)
tk.Button(button_frame, text="🚀 Analyze", width=13, command=analyze).grid(row=0, column=2, padx=3, pady=4)
tk.Button(button_frame, text="📋 Copy ATS Keywords", width=20, command=copy_ats_keywords).grid(row=0, column=3, padx=3, pady=4)
tk.Button(button_frame, text="📋 Copy Report", width=15, command=copy_report).grid(row=0, column=4, padx=3, pady=4)
tk.Button(button_frame, text="✨ Improve Resume", width=17, command=show_improve_resume_popup).grid(row=0, column=5, padx=3, pady=4)
tk.Button(button_frame, text="📤 Export Report", width=16, command=export_report).grid(row=0, column=6, padx=3, pady=4)
tk.Button(button_frame, text="📂 Reports", width=12, command=open_reports_folder).grid(row=0, column=7, padx=3, pady=4)
tk.Button(button_frame, text="🧹 Clear", width=10, command=clear_output).grid(row=0, column=8, padx=3, pady=4)

resume_label = tk.Label(app, text="Resume: Not loaded")
resume_label.pack()

job_label = tk.Label(app, text="Job Description: Not loaded")
job_label.pack()

status_label = tk.Label(app, text="Status: Ready", font=("Arial", 10, "bold"))
status_label.pack(pady=4)

progress_bar = ttk.Progressbar(app, orient="horizontal", length=800, mode="determinate", style="score.Horizontal.TProgressbar")
progress_bar.pack(pady=4)

output_box = scrolledtext.ScrolledText(app, width=140, height=26)
output_box.pack(pady=8)

footer_label = tk.Label(app, text="CareerCopilot v1.0 | Built by Domonique Robinson", font=("Arial", 9))
footer_label.pack(pady=4)

app.mainloop()